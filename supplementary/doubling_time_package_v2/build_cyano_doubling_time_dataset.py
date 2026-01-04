#!/usr/bin/env python3
"""
Build machine-readable cyanobacterial doubling time datasets.

This package was designed for a typical ML/DS repo workflow:
- Deterministic parsing of specific primary-source artifacts (PDF, DOCX)
- Tidy CSV outputs suitable for manuscript tables and downstream analysis
- Minimal external dependencies (pandas/numpy/matplotlib/pdfplumber/python-docx)

USAGE (recommended)
-------------------
If you ONLY want to inspect the already-built datasets shipped with this package, you do NOT
need to run this script. Instead, run:

    python analyze_cyano_doubling_times.py

If you want to REBUILD the datasets from primary sources, you have two options:

A) Provide local paths to the source files:
    python build_cyano_doubling_time_dataset.py \
      --yu2015_pdf /path/to/srep08132.pdf \
      --lab_docx /path/to/Cyano\ doubling\ time-122225.docx \
      --out_dir doubling_time_outputs

B) Let the script download the open-access Yu et al. (2015) PDF automatically:
    python build_cyano_doubling_time_dataset.py \
      --download_yu2015 \
      --lab_docx /path/to/Cyano\ doubling\ time-122225.docx \
      --out_dir doubling_time_outputs

NOTES
-----
- The Yu et al. (2015) Scientific Reports article is open access. The canonical PDF is:
  https://www.nature.com/articles/srep08132.pdf
- PDF text extraction can be brittle across different PDF renderings. This parser is
  tuned to the Nature-hosted PDF above.
- The lab DOCX is parsed from the first table and assumes the header row layout in the
  provided document.

Outputs (written to --out_dir)
------------------------------
- cyano_doubling_times_yu2015_srep08132_table1_long.csv
- cyano_doubling_times_yu2015_table1_strain_summary.csv
- cyano_doubling_times_lab_table2_ranges.csv
- cyano_doubling_times_compiled_best.csv
- cyano_doubling_times_compiled_best_plus_pcc11801.csv
- cyano_growth_throughput_10_doublings.csv
- yu2015_table1_doubling_time_heatmap.png
- README_doubling_time_outputs.md

References (for manuscript)
---------------------------
- Yu, J. et al. Synechococcus elongatus UTEX 2973, a fast growing cyanobacterial chassis for biosynthesis using light and CO2.
  Scientific Reports 5, 8132 (2015). doi:10.1038/srep08132.
- BioNumbers BNID 112485 (shortest UTEX 2973 doubling time, 1.9 h; derived from Yu et al. text).
- Jaiswal, D. et al. Genome Features and Biochemical Characteristics of a Rapidly Growing Cyanobacterium Synechococcus elongatus PCC 11801
  Isolated from an Urban Freshwater Lake. mBio (2018). (Doubling time 2.3 h under ambient CO2 conditions).

"""

from __future__ import annotations

import argparse
import math
import os
import re
import sys
import urllib.request
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

import matplotlib.pyplot as plt
import pdfplumber
from docx import Document


DEFAULT_YU2015_PDF_URL = "https://www.nature.com/articles/srep08132.pdf"


def _eprint(msg: str) -> None:
    print(msg, file=sys.stderr)


def download_file(url: str, dst_path: Path) -> None:
    """Download URL to dst_path (overwrites if exists). Uses only stdlib."""
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    _eprint(f"[download] {url} -> {dst_path}")
    with urllib.request.urlopen(url) as resp:
        data = resp.read()
    dst_path.write_bytes(data)
    _eprint(f"[download] wrote {dst_path.stat().st_size:,} bytes")


def resolve_existing_file(path_str: str) -> Optional[Path]:
    p = Path(path_str).expanduser()
    if p.is_file():
        return p
    return None


def search_for_file(patterns: List[str], root: Path, max_hits: int = 10) -> List[Path]:
    hits: List[Path] = []
    for pat in patterns:
        for p in root.rglob(pat):
            if p.is_file():
                hits.append(p)
                if len(hits) >= max_hits:
                    return hits
    return hits


def ensure_yu2015_pdf(
    yu2015_pdf: str,
    out_dir: Path,
    download_yu2015: bool,
    yu2015_url: str,
) -> Path:
    """
    Ensure we have a local Yu et al. 2015 PDF path.
    Strategy:
      1) Use --yu2015_pdf if it exists.
      2) Search cwd recursively for common filename patterns.
      3) If --download_yu2015, download from --yu2015_url to out_dir.
      4) Otherwise, raise with a helpful message.
    """
    p = resolve_existing_file(yu2015_pdf)
    if p is not None:
        return p

    # Common names users might have in repos
    patterns = [
        "srep08132.pdf",
        "*srep08132*.pdf",
        "*UTEX*2973*.pdf",
        "*Synechococcus*2973*.pdf",
        "*fast*grow*2015*.pdf",
        "*fast*grow*cyanobacter*pdf",
    ]
    hits = search_for_file(patterns, root=Path.cwd())
    if hits:
        _eprint(f"[info] --yu2015_pdf not found ('{yu2015_pdf}'). Using discovered PDF: {hits[0]}")
        return hits[0]

    if download_yu2015:
        dst = out_dir / "yu2015_srep08132.pdf"
        download_file(yu2015_url, dst)
        return dst

    raise FileNotFoundError(
        "Yu et al. 2015 PDF not found.\n"
        f"  Provided --yu2015_pdf: {yu2015_pdf}\n"
        "  Also searched the current directory tree for typical filenames but found nothing.\n\n"
        "Fix options:\n"
        "  (1) Re-run with the correct local path, e.g. --yu2015_pdf /path/to/srep08132.pdf\n"
        "  (2) Re-run with --download_yu2015 to download the open-access PDF automatically.\n"
        f"      Default URL: {yu2015_url}\n"
    )


def parse_value_token(tok: str) -> Dict[str, object]:
    """Parse tokens like '2.160.2' (PDF extraction of '2.1±0.2'), '-' (ND) and 'NG'."""
    tok = tok.strip()
    if tok in {"-", ""}:
        return {"value_h": np.nan, "sd_h": np.nan, "qualifier": "ND"}
    if tok.upper() == "NG":
        return {"value_h": np.nan, "sd_h": np.nan, "qualifier": "NG"}

    # In the Nature-hosted PDF, '±' may be extracted as the character '6' between mean and SD.
    m = re.match(r"^(\d+(?:\.\d+)?)6(\d+(?:\.\d+)?)$", tok)
    if m:
        return {"value_h": float(m.group(1)), "sd_h": float(m.group(2)), "qualifier": ""}

    try:
        return {"value_h": float(tok), "sd_h": np.nan, "qualifier": ""}
    except ValueError:
        return {"value_h": np.nan, "sd_h": np.nan, "qualifier": f"UNPARSEABLE:{tok}"}


def parse_yu2015_table1(pdf_path: Path) -> pd.DataFrame:
    """Extract Table 1 (doubling times) from the Yu et al. 2015 Sci Rep PDF (Nature-hosted)."""
    with pdfplumber.open(str(pdf_path)) as pdf:
        # Table 1 appears on page 2 in the Nature PDF (0-index: 1)
        page = pdf.pages[1]
        txt = page.extract_text() or ""

    lines = txt.splitlines()

    # Locate the table header by the 'Temp' line and stop at 'Means'
    start = None
    means_i = None
    for i, l in enumerate(lines):
        if l.startswith("Temp") and start is None:
            start = i
        if l.strip().startswith("Means") and means_i is None:
            means_i = i
    if start is None or means_i is None:
        raise RuntimeError(
            "Could not locate Table 1 header (Temp/Means lines not found).\n"
            "This usually means the PDF text extraction differs from the expected Nature-hosted PDF.\n"
            f"PDF: {pdf_path}"
        )

    # Parse light intensity line (contains 8 columns)
    light_line = lines[start + 3]
    light_vals = [int(x) for x in light_line.split()[1:]]

    # Column mapping inferred from the table's multi-level header:
    # 2 columns at 41°C (3% CO2), 4 columns at 38°C (3% CO2), 2 columns at 30°C (Air then 3% CO2)
    temps = [41] * 2 + [38] * 4 + [30] * 2
    co2s = ["3%"] * 6 + ["Air (0.04%)", "3%"]
    if not (len(temps) == len(co2s) == len(light_vals) == 8):
        raise RuntimeError("Unexpected number of columns parsed from Table 1 header.")

    rows: List[Dict[str, object]] = []
    for l in lines[start + 4 : means_i]:
        if not (l.startswith("UTEX") or l.startswith("PCC")):
            continue
        toks = l.split()
        strain = toks[0]
        toks_clean = [t for t in toks[1:] if not t.lower().startswith("doubling")]
        if len(toks_clean) != 8:
            raise RuntimeError(f"Unexpected token count for {strain}: {len(toks_clean)} tokens: {toks_clean}")
        for j, tok in enumerate(toks_clean):
            parsed = parse_value_token(tok)
            rows.append(
                {
                    "strain": strain,
                    "condition_idx": j + 1,
                    "temp_C": temps[j],
                    "co2": co2s[j],
                    "light_umol_photons_m2_s": light_vals[j],
                    "doubling_time_h": parsed["value_h"],
                    "doubling_time_sd_h": parsed["sd_h"],
                    "qualifier": parsed["qualifier"],
                    "source": "Yu et al. 2015 Sci Rep 5:8132 (srep08132) Table 1",
                }
            )

    return pd.DataFrame(rows)


def parse_range(s: str) -> Tuple[float, float, str]:
    s = (s or "").strip()
    if s == "":
        return (np.nan, np.nan, "")
    if s.startswith(">") and len(s) > 1:
        try:
            val = float(s[1:])
            return (val, np.nan, ">")
        except ValueError:
            return (np.nan, np.nan, "")
    if s.startswith("~") and len(s) > 1:
        try:
            val = float(s[1:])
            return (val, val, "~")
        except ValueError:
            return (np.nan, np.nan, "")
    if "-" in s:
        parts = s.split("-")
        if len(parts) == 2:
            try:
                return (float(parts[0]), float(parts[1]), "range")
            except ValueError:
                return (np.nan, np.nan, "")
    try:
        val = float(s)
        return (val, val, "")
    except ValueError:
        return (np.nan, np.nan, "")


def parse_lab_docx_table2(docx_path: Path) -> pd.DataFrame:
    doc = Document(str(docx_path))
    if len(doc.tables) < 1:
        raise RuntimeError("No tables found in DOCX.")
    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in r.cells] for r in table.rows]

    # The provided DOCX has title rows before the header; we keep the same logic as v1
    headers = rows[1]
    data_rows = rows[2:]
    df = pd.DataFrame(data_rows, columns=headers)

    mins, maxs, quals = [], [], []
    for s in df["Doubling Time (hr)"]:
        mn, mx, q = parse_range(str(s))
        mins.append(mn)
        maxs.append(mx)
        quals.append(q)

    df_out = pd.DataFrame(
        {
            "strain": df["Strain"],
            "temp_C": pd.to_numeric(df["Growth Temp (C°)"], errors="coerce"),
            "co2": np.nan,
            "light_umol_photons_m2_s": np.nan,
            "doubling_time_min_h": mins,
            "doubling_time_max_h": maxs,
            "doubling_time_mid_h": [
                (a + b) / 2 if (np.isfinite(a) and np.isfinite(b)) else np.nan for a, b in zip(mins, maxs)
            ],
            "qualifier": quals,
            "metabolism": df["Metabolism"],
            "notes": df["Notes"],
            "reference": df["References"],
            "source": "Cyano doubling time-122225.docx Table 2 (adapted from Berla et al., 2013)",
        }
    )
    return df_out


def make_heatmap_table1(df_table1: pd.DataFrame, out_png: Path) -> None:
    # Build a strain x condition matrix for visualization
    strains = ["UTEX2973", "PCC7942", "PCC7002", "PCC6301", "PCC6803"]

    # infer condition labels from the first strain rows (assumes all conditions present)
    conds = (
        df_table1.sort_values("condition_idx")
        .drop_duplicates("condition_idx")
        .sort_values("condition_idx")[["condition_idx", "temp_C", "co2", "light_umol_photons_m2_s"]]
    )
    cond_labels = [
        f"T{int(r.temp_C)}_{str(r.co2).replace(' ', '')}_L{int(r.light_umol_photons_m2_s)}" for r in conds.itertuples()
    ]

    mat = np.full((len(strains), len(conds)), np.nan)
    qual_mat = np.full((len(strains), len(conds)), "", dtype=object)

    for i, strain in enumerate(strains):
        sub = df_table1[df_table1["strain"] == strain].sort_values("condition_idx")
        for r in sub.itertuples():
            j = int(r.condition_idx) - 1
            mat[i, j] = r.doubling_time_h
            qual_mat[i, j] = r.qualifier or ""

    fig, ax = plt.subplots(figsize=(11, 3.2))
    im = ax.imshow(mat, aspect="auto")
    ax.set_yticks(range(len(strains)))
    ax.set_yticklabels(strains)
    ax.set_xticks(range(len(conds)))
    ax.set_xticklabels(cond_labels, rotation=45, ha="right", fontsize=8)
    ax.set_title("Doubling time (h) across conditions (Yu et al. Sci Rep 2015 Table 1)")
    cbar = fig.colorbar(im, ax=ax)
    cbar.set_label("Doubling time (h)")

    # annotate values / qualifiers
    for i in range(len(strains)):
        for j in range(len(conds)):
            v = mat[i, j]
            if np.isfinite(v):
                ax.text(j, i, f"{v:g}", ha="center", va="center", fontsize=7)
            else:
                q = qual_mat[i, j]
                if q == "NG":
                    ax.text(j, i, "NG", ha="center", va="center", fontsize=7)
                elif q == "ND":
                    ax.text(j, i, "–", ha="center", va="center", fontsize=7)

    fig.tight_layout()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(str(out_png), dpi=200)
    plt.close(fig)


def compute_mu_per_h(doubling_time_h: float) -> float:
    if not np.isfinite(doubling_time_h) or doubling_time_h <= 0:
        return np.nan
    return math.log(2.0) / doubling_time_h


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--yu2015_pdf", default="S. elongatus, fastest-growing2015.pdf", help="Path to Yu et al. 2015 PDF containing Table 1")
    ap.add_argument("--lab_docx", default="Cyano doubling time-122225.docx", help="Path to lab doubling time DOCX")
    ap.add_argument("--out_dir", default="doubling_time_outputs", help="Output directory")
    ap.add_argument("--download_yu2015", action="store_true", help="Download Yu et al. PDF if not found locally (open access)")
    ap.add_argument("--yu2015_url", default=DEFAULT_YU2015_PDF_URL, help="URL for Yu et al. (2015) PDF download")
    ap.add_argument("--skip_yu2015", action="store_true", help="Skip parsing Yu et al. Table 1 (build only from lab DOCX)")
    ap.add_argument("--skip_lab", action="store_true", help="Skip parsing lab DOCX table (build only from Yu et al.)")
    ap.add_argument("--include_pcc11801", action="store_true", help="Include PCC 11801 best-case (2.3 h, ambient CO2) in compiled_best_plus_pcc11801.csv")
    args = ap.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # -------------------------
    # Yu et al. 2015 Table 1
    # -------------------------
    df_table1: Optional[pd.DataFrame] = None
    if not args.skip_yu2015:
        yu_pdf = ensure_yu2015_pdf(
            yu2015_pdf=args.yu2015_pdf,
            out_dir=out_dir,
            download_yu2015=args.download_yu2015,
            yu2015_url=args.yu2015_url,
        )
        _eprint(f"[info] parsing Yu et al. Table 1 from: {yu_pdf}")
        df_table1 = parse_yu2015_table1(yu_pdf)
        df_table1.to_csv(out_dir / "cyano_doubling_times_yu2015_srep08132_table1_long.csv", index=False)

        # strain-level summary
        summary = (
            df_table1.dropna(subset=["doubling_time_h"])
            .groupby("strain")
            .agg(
                n_conditions=("doubling_time_h", "count"),
                min_doubling_time_h=("doubling_time_h", "min"),
                max_doubling_time_h=("doubling_time_h", "max"),
                median_doubling_time_h=("doubling_time_h", "median"),
            )
            .reset_index()
        )
        summary.to_csv(out_dir / "cyano_doubling_times_yu2015_table1_strain_summary.csv", index=False)

        # heatmap
        make_heatmap_table1(df_table1, out_dir / "yu2015_table1_doubling_time_heatmap.png")

    # -------------------------
    # Lab DOCX table
    # -------------------------
    df_lab: Optional[pd.DataFrame] = None
    if not args.skip_lab:
        lab_docx_path = resolve_existing_file(args.lab_docx)
        if lab_docx_path is None:
            # try to discover it in the current tree
            hits = search_for_file(["*doubling*time*.docx", "*Cyano*double*.docx", "*.docx"], root=Path.cwd(), max_hits=50)
            hits = [h for h in hits if "doubling" in h.name.lower() and h.suffix.lower() == ".docx"]
            if hits:
                lab_docx_path = hits[0]
                _eprint(f"[info] --lab_docx not found ('{args.lab_docx}'). Using discovered DOCX: {lab_docx_path}")
        if lab_docx_path is None:
            raise FileNotFoundError(
                "Lab DOCX not found.\n"
                f"  Provided --lab_docx: {args.lab_docx}\n"
                "Fix: re-run with the correct path to Cyano doubling time-122225.docx (or use --skip_lab)."
            )
        _eprint(f"[info] parsing lab DOCX table from: {lab_docx_path}")
        df_lab = parse_lab_docx_table2(lab_docx_path)
        df_lab.to_csv(out_dir / "cyano_doubling_times_lab_table2_ranges.csv", index=False)

    # -------------------------
    # Compiled tables
    # -------------------------
    compiled_rows: List[Dict[str, object]] = []

    # Best-case per strain from Yu table
    if df_table1 is not None and not df_table1.empty:
        best_table1 = (
            df_table1.dropna(subset=["doubling_time_h"])
            .sort_values(["strain", "doubling_time_h"])
            .groupby("strain", as_index=False)
            .first()
        )
        for r in best_table1.itertuples():
            compiled_rows.append(
                dict(
                    strain=r.strain,
                    temp_C=r.temp_C,
                    co2=r.co2,
                    light_umol_photons_m2_s=r.light_umol_photons_m2_s,
                    doubling_time_min_h=r.doubling_time_h,
                    doubling_time_max_h=r.doubling_time_h,
                    doubling_time_sd_h=r.doubling_time_sd_h,
                    qualifier="",
                    reference="Yu et al., 2015 (Sci Rep) Table 1",
                    source=r.source,
                    data_type="condition-specific",
                )
            )

    # Add lab summary rows
    if df_lab is not None and not df_lab.empty:
        for r in df_lab.itertuples():
            compiled_rows.append(
                dict(
                    strain=r.strain,
                    temp_C=r.temp_C,
                    co2=r.co2,
                    light_umol_photons_m2_s=r.light_umol_photons_m2_s,
                    doubling_time_min_h=r.doubling_time_min_h,
                    doubling_time_max_h=r.doubling_time_max_h,
                    doubling_time_sd_h=np.nan,
                    qualifier=r.qualifier,
                    reference=r.reference,
                    source=r.source,
                    data_type="approx/range",
                )
            )

    compiled = pd.DataFrame(compiled_rows)

    # Add BioNumbers shortest doubling time for UTEX2973 (1.9 h) if Yu table was included
    if df_table1 is not None:
        compiled = pd.concat(
            [
                compiled,
                pd.DataFrame(
                    [
                        dict(
                            strain="UTEX2973",
                            temp_C=41,
                            co2="3%",
                            light_umol_photons_m2_s=500,
                            doubling_time_min_h=1.9,
                            doubling_time_max_h=1.9,
                            doubling_time_sd_h=np.nan,
                            qualifier="",
                            reference="BioNumbers BNID 112485 (derived from Yu et al. 2015 text)",
                            source="BioNumbers BNID 112485 (shortest doubling time)",
                            data_type="condition-specific",
                        )
                    ]
                ),
            ],
            ignore_index=True,
        )

    if not compiled.empty:
        compiled["mu_per_h"] = compiled["doubling_time_min_h"].apply(compute_mu_per_h)
        compiled.to_csv(out_dir / "cyano_doubling_times_compiled_best.csv", index=False)

    # Add PCC11801 row (ambient CO2 best-case) to a second CSV to keep provenance explicit
    compiled_plus = compiled.copy()
    if args.include_pcc11801:
        compiled_plus = pd.concat(
            [
                compiled_plus,
                pd.DataFrame(
                    [
                        dict(
                            strain="PCC11801",
                            temp_C=np.nan,
                            co2="Air (0.04%)",
                            light_umol_photons_m2_s=np.nan,
                            doubling_time_min_h=2.3,
                            doubling_time_max_h=2.3,
                            doubling_time_sd_h=np.nan,
                            qualifier="",
                            reference="Jaiswal et al., 2018 (mBio) (doubling time under ambient CO2)",
                            source="mBio 2018 PCC11801 physiology/genome paper",
                            data_type="condition-specific",
                            mu_per_h=compute_mu_per_h(2.3),
                        )
                    ]
                ),
            ],
            ignore_index=True,
        )
    compiled_plus.to_csv(out_dir / "cyano_doubling_times_compiled_best_plus_pcc11801.csv", index=False)

    # Derived throughput summary (time to 10 doublings)
    if not compiled_plus.empty:
        thr = compiled_plus.copy()
        thr["time_for_10_doublings_h"] = thr["doubling_time_min_h"] * 10.0
        thr = thr[
            [
                "strain",
                "doubling_time_min_h",
                "mu_per_h",
                "time_for_10_doublings_h",
                "temp_C",
                "co2",
                "light_umol_photons_m2_s",
                "reference",
                "data_type",
            ]
        ].sort_values("time_for_10_doublings_h")
        thr.to_csv(out_dir / "cyano_growth_throughput_10_doublings.csv", index=False)

    # README
    with open(out_dir / "README_doubling_time_outputs.md", "w") as f:
        f.write("# Cyanobacterial doubling time datasets (compiled)\n\n")
        f.write("Generated by build_cyano_doubling_time_dataset.py\n\n")
        f.write("## Sources\n")
        if df_table1 is not None:
            f.write("- Yu et al. 2015 Sci Rep 5:8132 (srep08132), Table 1\n")
            f.write("- BioNumbers BNID 112485 (shortest UTEX 2973 doubling time)\n")
        if df_lab is not None:
            f.write("- Lab DOCX: Cyano doubling time-122225.docx, Table 2 (adapted from Berla et al., 2013)\n")
        if args.include_pcc11801:
            f.write("- PCC 11801 doubling time (ambient CO2) from Jaiswal et al. 2018 (mBio)\n")

    _eprint(f"[done] outputs written to: {out_dir.resolve()}")


if __name__ == "__main__":
    main()
